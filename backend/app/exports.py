from __future__ import annotations

import io
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Mm as DocxMm
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


STATUS_LABELS = {
    "unresearched": "尚未研究", "ai": "AI 初稿", "edited": "人工整理",
    "evidenced": "已有证据", "verified": "已核验", "stale": "需要更新",
}


def ordered_nodes(state: dict[str, Any], only_expanded: bool = False) -> list[tuple[str, dict[str, Any], int]]:
    nodes = state.get("nodes") or {}
    root_id = state.get("rootId") or next(iter(nodes), None)
    expanded = set(state.get("expanded") or [])
    result: list[tuple[str, dict[str, Any], int]] = []
    seen: set[str] = set()

    def walk(node_id: str, depth: int) -> None:
        if node_id in seen or node_id not in nodes:
            return
        seen.add(node_id)
        result.append((node_id, nodes[node_id], depth))
        if only_expanded and node_id not in expanded:
            return
        for child_id in nodes[node_id].get("children") or []:
            walk(child_id, depth + 1)

    if root_id:
        walk(root_id, 0)
    for node_id in nodes:
        if node_id not in seen:
            walk(node_id, 1)
    return result


def _sources(state: dict[str, Any], node_id: str) -> list[dict[str, Any]]:
    return (state.get("evidenceData") or {}).get(node_id) or []


def _dims(state: dict[str, Any], node_id: str) -> list[list[str]]:
    return ((state.get("researchData") or {}).get(node_id) or {}).get("dims") or []


def build_docx(state: dict[str, Any], title: str, include_evidence: bool, include_tasks: bool, only_expanded: bool, document_tables: list[dict[str, Any]]) -> bytes:
    doc = Document()
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    doc.add_heading(title, 0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}｜Industry Atlas 精准导出")
    ordered = ordered_nodes(state, only_expanded)
    doc.add_heading("一、产业结构", level=1)
    for _, node, depth in ordered:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = DocxMm(depth * 7)
        p.add_run(node.get("title") or "未命名节点").bold = True
        p.add_run(f" — {node.get('desc') or node.get('summary') or ''}")
    doc.add_page_break()
    doc.add_heading("二、节点研究", level=1)
    tasks = state.get("researchTasks") or []
    for index, (node_id, node, _) in enumerate(ordered):
        doc.add_heading(node.get("title") or "未命名节点", level=2)
        meta = doc.add_table(rows=3, cols=2)
        meta.style = "Light Shading Accent 1"
        meta.cell(0, 0).text, meta.cell(0, 1).text = "产业位置", str(node.get("category") or "未标注")
        meta.cell(1, 0).text, meta.cell(1, 1).text = "研究状态", STATUS_LABELS.get(node.get("status"), str(node.get("status") or "未标注"))
        meta.cell(2, 0).text, meta.cell(2, 1).text = "更新时间", str(node.get("updatedAt") or "未标注")
        doc.add_heading("核心解释", level=3)
        doc.add_paragraph(str(node.get("summary") or "暂无"))
        doc.add_heading("为什么重要", level=3)
        doc.add_paragraph(str(node.get("why") or "暂无"))
        dims = _dims(state, node_id)
        if dims:
            doc.add_heading("研究维度", level=3)
            for label, value, *_ in dims:
                doc.add_paragraph(f"{label}：{value}", style="List Bullet")
        if include_evidence:
            sources = _sources(state, node_id)
            doc.add_heading("证据与原文引用", level=3)
            if not sources:
                doc.add_paragraph("暂无证据资料")
            for number, source in enumerate(sources, 1):
                p = doc.add_paragraph()
                p.add_run(f"[{number}] {source.get('title') or '未命名资料'}").bold = True
                p.add_run(f"｜{source.get('location') or '位置未标注'}｜{'已人工核验' if source.get('verified') else '待人工核验'}")
                if source.get("quote"):
                    doc.add_paragraph(f"原文摘录：{source['quote']}")
        if include_tasks:
            pending = [task for task in tasks if task.get("nodeId") == node_id and task.get("status") != "done"]
            if pending:
                doc.add_heading("未完成任务", level=3)
                for task in pending:
                    doc.add_paragraph(f"{task.get('title')}（{task.get('status')} / {task.get('priority')}）", style="List Bullet")
        if index < len(ordered) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
    companies = state.get("companyData") or []
    if companies:
        doc.add_page_break()
        doc.add_heading("三、公司与产业映射", level=1)
        for company in companies:
            doc.add_heading(str(company.get("name") or "未命名公司"), level=2)
            doc.add_paragraph(f"报告期：{company.get('reportPeriod') or '未识别'}")
            doc.add_paragraph(str(company.get("summary") or "暂无公司摘要"))
            for period in company.get("periods") or []:
                doc.add_paragraph(f"{period.get('period')}｜{len(period.get('documents') or [])} 份年报｜{len(period.get('findingIds') or [])} 条发现｜{period.get('summary') or ''}", style="List Bullet")
            for mapping in company.get("mappings") or []:
                node_name = str((state.get("nodes") or {}).get(mapping.get("nodeId"), {}).get("title") or "已删除节点")
                doc.add_paragraph(f"{node_name}｜{mapping.get('status')}｜匹配分 {mapping.get('score')}｜{mapping.get('reason') or ''}", style="List Bullet")
    if document_tables:
        doc.add_page_break()
        doc.add_heading("四、年报表格附录", level=1)
        for table_info in document_tables:
            doc.add_heading(f"{table_info['filename']}｜第 {table_info['page_number']} 页｜表格 {table_info['table_number']}", level=2)
            data = table_info.get("data") or []
            if not data:
                continue
            table = doc.add_table(rows=len(data), cols=max(len(row) for row in data))
            table.style = "Table Grid"
            for row_index, row in enumerate(data):
                for column_index, value in enumerate(row):
                    table.cell(row_index, column_index).text = str(value or "")
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _sheet(ws, headers: list[str]) -> None:
    ws.append(headers)
    fill = PatternFill("solid", fgColor="6255D3")
    for cell in ws[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = fill
        cell.alignment = Alignment(vertical="center")
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def build_xlsx(state: dict[str, Any], title: str, include_evidence: bool, include_tasks: bool, only_expanded: bool, document_tables: list[dict[str, Any]]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "节点"
    _sheet(ws, ["节点ID", "节点名称", "层级", "产业位置", "状态", "一句话解释", "为什么重要", "更新时间"])
    allowed = set()
    for node_id, node, depth in ordered_nodes(state, only_expanded):
        allowed.add(node_id)
        ws.append([node_id, node.get("title"), depth, node.get("category"), STATUS_LABELS.get(node.get("status"), node.get("status")), node.get("summary"), node.get("why"), node.get("updatedAt")])
    relations = wb.create_sheet("关系")
    _sheet(relations, ["来源节点ID", "来源节点", "目标节点ID", "目标节点", "关系类型", "关系说明"])
    nodes = state.get("nodes") or {}
    for edge in state.get("edges") or []:
        if edge.get("source") in allowed and edge.get("target") in allowed:
            relations.append([edge.get("source"), (nodes.get(edge.get("source")) or {}).get("title"), edge.get("target"), (nodes.get(edge.get("target")) or {}).get("title"), edge.get("type"), edge.get("label") or edge.get("reason")])
    evidence = wb.create_sheet("证据")
    _sheet(evidence, ["节点ID", "节点名称", "资料类型", "资料标题", "日期", "页码/位置", "原文摘录", "人工核验", "原文匹配", "链接", "文档ID", "提取任务ID"])
    if include_evidence:
        for node_id in allowed:
            for source in _sources(state, node_id):
                evidence.append([node_id, (nodes.get(node_id) or {}).get("title"), source.get("type"), source.get("title"), source.get("date"), source.get("location"), source.get("quote"), "是" if source.get("verified") else "否", "是" if source.get("citationMatched") else "否", source.get("url"), source.get("documentId"), source.get("extractionId")])
    tasks_ws = wb.create_sheet("任务")
    _sheet(tasks_ws, ["任务ID", "节点ID", "节点名称", "任务", "类型", "优先级", "状态", "备注"])
    if include_tasks:
        for task in state.get("researchTasks") or []:
            if task.get("nodeId") in allowed:
                tasks_ws.append([task.get("id"), task.get("nodeId"), (nodes.get(task.get("nodeId")) or {}).get("title"), task.get("title"), task.get("type"), task.get("priority"), task.get("status"), task.get("note")])
    companies_ws = wb.create_sheet("公司")
    _sheet(companies_ws, ["公司ID", "公司名称", "报告期", "公司摘要", "年报数量", "结构化发现数量", "更新时间"])
    mappings_ws = wb.create_sheet("公司映射")
    _sheet(mappings_ws, ["公司ID", "公司名称", "节点ID", "节点名称", "状态", "匹配分", "映射理由", "审核时间"])
    periods_ws = wb.create_sheet("公司期间")
    _sheet(periods_ws, ["公司ID", "公司名称", "报告期", "期间摘要", "年报数量", "发现数量"])
    metrics_ws = wb.create_sheet("公司指标")
    _sheet(metrics_ws, ["公司ID", "公司名称", "报告期", "指标类型", "指标名称", "指标值", "原文", "页码", "原文匹配"])
    for company in state.get("companyData") or []:
        companies_ws.append([company.get("id"), company.get("name"), company.get("reportPeriod"), company.get("summary"), len(company.get("documents") or []), len(company.get("findings") or []), company.get("updatedAt")])
        for mapping in company.get("mappings") or []:
            mappings_ws.append([company.get("id"), company.get("name"), mapping.get("nodeId"), (nodes.get(mapping.get("nodeId")) or {}).get("title"), mapping.get("status"), mapping.get("score"), mapping.get("reason"), mapping.get("reviewedAt")])
        for period in company.get("periods") or []:
            periods_ws.append([company.get("id"), company.get("name"), period.get("period"), period.get("summary"), len(period.get("documents") or []), len(period.get("findingIds") or [])])
        for finding in company.get("findings") or []:
            if finding.get("category") in {"financial", "capacity"}:
                metrics_ws.append([company.get("id"), company.get("name"), finding.get("reportPeriod") or company.get("reportPeriod"), finding.get("category_label") or finding.get("category"), finding.get("title"), finding.get("value"), finding.get("quote"), "、".join(str(page) for page in (finding.get("matched_pages") or finding.get("page_numbers") or [])), "是" if finding.get("citation_status") == "matched" else "否"])
    table_ws = wb.create_sheet("年报表格")
    max_columns = max((len(row) for table in document_tables for row in (table.get("data") or [])), default=0)
    _sheet(table_ws, ["文件", "页码", "表格序号", "行号", *[f"字段{i}" for i in range(1, max_columns + 1)]])
    for table in document_tables:
        for row_number, row in enumerate(table.get("data") or [], start=1):
            table_ws.append([table.get("filename"), table.get("page_number"), table.get("table_number"), row_number, *row])
    for sheet in wb.worksheets:
        sheet.auto_filter.ref = sheet.dimensions
        for column in sheet.columns:
            letter = column[0].column_letter
            width = min(60, max(10, max(len(str(cell.value or "")) for cell in column) + 2))
            sheet.column_dimensions[letter].width = width
            for cell in column:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


def _pdf_font() -> str:
    if "IndustryAtlasCJK" in pdfmetrics.getRegisteredFontNames():
        return "IndustryAtlasCJK"
    candidates = [Path("C:/Windows/Fonts/msyh.ttc"), Path("C:/Windows/Fonts/simsun.ttc")]
    for path in candidates:
        if path.exists():
            pdfmetrics.registerFont(TTFont("IndustryAtlasCJK", str(path), subfontIndex=0))
            pdfmetrics.registerFontFamily("IndustryAtlasCJK", normal="IndustryAtlasCJK", bold="IndustryAtlasCJK", italic="IndustryAtlasCJK", boldItalic="IndustryAtlasCJK")
            return "IndustryAtlasCJK"
    return "Helvetica"


def build_pdf(state: dict[str, Any], title: str, include_evidence: bool, include_tasks: bool, only_expanded: bool, document_tables: list[dict[str, Any]]) -> bytes:
    output = io.BytesIO()
    font = _pdf_font()
    styles = getSampleStyleSheet()
    body = ParagraphStyle("CJKBody", parent=styles["BodyText"], fontName=font, fontSize=9, leading=15, textColor=colors.HexColor("#39465B"))
    h1 = ParagraphStyle("CJKH1", parent=body, fontSize=20, leading=28, alignment=TA_CENTER, textColor=colors.HexColor("#302A62"), spaceAfter=12)
    h2 = ParagraphStyle("CJKH2", parent=body, fontSize=14, leading=21, textColor=colors.HexColor("#4E43AA"), spaceBefore=12, spaceAfter=6)
    h3 = ParagraphStyle("CJKH3", parent=body, fontSize=11, leading=17, textColor=colors.HexColor("#273349"), spaceBefore=8, spaceAfter=4)
    doc = SimpleDocTemplate(output, pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=18*mm, bottomMargin=18*mm, title=title)
    story = [Paragraph(escape(title), h1), Paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}｜Industry Atlas 精准导出", body), Spacer(1, 8*mm), Paragraph("一、产业结构", h2)]
    ordered = ordered_nodes(state, only_expanded)
    for _, node, depth in ordered:
        story.append(Paragraph(f"{'　' * depth}• <b>{escape(str(node.get('title') or '未命名节点'))}</b> — {escape(str(node.get('desc') or node.get('summary') or ''))}", body))
    story.extend([PageBreak(), Paragraph("二、节点研究", h2)])
    tasks = state.get("researchTasks") or []
    for node_id, node, _ in ordered:
        story.append(Paragraph(escape(str(node.get("title") or "未命名节点")), h2))
        meta = [["产业位置", str(node.get("category") or "未标注")], ["研究状态", STATUS_LABELS.get(node.get("status"), str(node.get("status") or "未标注"))]]
        table = Table([[Paragraph(escape(a), body), Paragraph(escape(b), body)] for a, b in meta], colWidths=[28*mm, 125*mm])
        table.setStyle(TableStyle([("GRID", (0,0), (-1,-1), .3, colors.HexColor("#DDE2E9")), ("BACKGROUND", (0,0), (0,-1), colors.HexColor("#F1EFFF")), ("VALIGN", (0,0), (-1,-1), "TOP"), ("LEFTPADDING", (0,0), (-1,-1), 6), ("RIGHTPADDING", (0,0), (-1,-1), 6)]))
        story.extend([table, Paragraph("核心解释", h3), Paragraph(escape(str(node.get("summary") or "暂无")), body), Paragraph("为什么重要", h3), Paragraph(escape(str(node.get("why") or "暂无")), body)])
        for label, value, *_ in _dims(state, node_id):
            story.append(Paragraph(f"• <b>{escape(str(label))}</b>：{escape(str(value))}", body))
        if include_evidence:
            story.append(Paragraph("证据与原文引用", h3))
            sources = _sources(state, node_id)
            if not sources:
                story.append(Paragraph("暂无证据资料", body))
            for number, source in enumerate(sources, 1):
                story.append(Paragraph(f"[{number}] <b>{escape(str(source.get('title') or '未命名资料'))}</b>｜{escape(str(source.get('location') or '位置未标注'))}｜{'已人工核验' if source.get('verified') else '待人工核验'}", body))
                if source.get("quote"):
                    story.append(Paragraph(f"原文摘录：{escape(str(source['quote']))}", body))
        if include_tasks:
            pending = [task for task in tasks if task.get("nodeId") == node_id and task.get("status") != "done"]
            if pending:
                story.append(Paragraph("未完成任务", h3))
                for task in pending:
                    story.append(Paragraph(f"• {escape(str(task.get('title')))}（{escape(str(task.get('status')))} / {escape(str(task.get('priority'))) }）", body))
        story.append(Spacer(1, 5*mm))
    companies = state.get("companyData") or []
    if companies:
        story.extend([PageBreak(), Paragraph("三、公司与产业映射", h2)])
        nodes = state.get("nodes") or {}
        for company in companies:
            story.append(Paragraph(escape(str(company.get("name") or "未命名公司")), h2))
            story.append(Paragraph(f"报告期：{escape(str(company.get('reportPeriod') or '未识别'))}｜{escape(str(company.get('summary') or '暂无公司摘要'))}", body))
            for period in company.get("periods") or []:
                story.append(Paragraph(f"• {escape(str(period.get('period')))}｜{len(period.get('documents') or [])} 份年报｜{len(period.get('findingIds') or [])} 条发现｜{escape(str(period.get('summary') or ''))}", body))
            for mapping in company.get("mappings") or []:
                node_name = (nodes.get(mapping.get("nodeId")) or {}).get("title") or "已删除节点"
                story.append(Paragraph(f"• {escape(str(node_name))}｜{escape(str(mapping.get('status')))}｜匹配分 {mapping.get('score')}｜{escape(str(mapping.get('reason') or ''))}", body))
    if document_tables:
        story.extend([PageBreak(), Paragraph("四、年报表格附录", h2)])
        for table in document_tables:
            story.append(Paragraph(f"{escape(str(table.get('filename')))}｜第 {table.get('page_number')} 页｜表格 {table.get('table_number')}", h3))
            for row_number, row in enumerate(table.get("data") or [], start=1):
                story.append(Paragraph(f"{row_number}. {escape(' ｜ '.join(str(value or '') for value in row))}", body))
    doc.build(story)
    return output.getvalue()


def build_export(state: dict[str, Any], title: str, format_name: str, include_evidence: bool, include_tasks: bool, only_expanded: bool, document_tables: list[dict[str, Any]] | None = None) -> tuple[bytes, str, str]:
    tables = document_tables or []
    if format_name == "docx":
        return build_docx(state, title, include_evidence, include_tasks, only_expanded, tables), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
    if format_name == "xlsx":
        return build_xlsx(state, title, include_evidence, include_tasks, only_expanded, tables), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"
    if format_name == "pdf":
        return build_pdf(state, title, include_evidence, include_tasks, only_expanded, tables), "application/pdf", "pdf"
    raise ValueError("不支持的导出格式")
